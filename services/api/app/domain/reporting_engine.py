import csv
import hashlib
import io
import json
from datetime import datetime, timezone
from pathlib import Path
from xml.etree.ElementTree import Element, SubElement, tostring

from docx import Document
from openpyxl import Workbook
from openpyxl.styles import Font
from reportlab.lib.pagesizes import A4, landscape
from reportlab.pdfgen import canvas
from sqlalchemy import select
from sqlalchemy.orm import Session

from app.db.models import Plant, ReportArtifact, ReportRun, ShiftEntry, Ticket
from app.settings import get_settings


MIME_BY_FORMAT = {
    "csv": "text/csv",
    "json": "application/json",
    "xml": "application/xml",
    "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "pdf": "application/pdf",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}
SUPPORTED_REPORT_KINDS = {"tickets", "schichtbuch", "kombiniert"}
CLOSED_TICKET_STATUSES = {"CLOSED", "CANCELLED", "CANCELLED_WRONG_PLANT"}


def normalize_report_kind(value: str | None) -> str:
    kind = (value or "tickets").strip().lower()
    if kind not in SUPPORTED_REPORT_KINDS:
        return "tickets"
    return kind


def _iso(ts: datetime | None) -> str | None:
    if ts is None:
        return None
    return ts.isoformat()


def _ticket_rows(
    db: Session,
    *,
    plant_slug: str,
    from_dt: datetime | None,
    to_dt: datetime | None,
    department: str | None,
    ticket_id: int | None,
    limit: int,
) -> list[dict]:
    stmt = select(Ticket, Plant.slug).join(Plant, Plant.id == Ticket.plant_id).where(Plant.slug == plant_slug)
    if department:
        stmt = stmt.where(Ticket.department == department)
    if ticket_id is not None:
        stmt = stmt.where(Ticket.id == ticket_id)
    if from_dt is not None:
        stmt = stmt.where(Ticket.created_at >= from_dt)
    if to_dt is not None:
        stmt = stmt.where(Ticket.created_at <= to_dt)

    rows = db.execute(stmt.order_by(Ticket.created_at.desc()).limit(max(1, min(limit, 2000)))).all()
    return [
        {
            "row_type": "ticket",
            "ticket_id": t.id,
            "plant_slug": slug,
            "bereich": t.department,
            "ticket_typ": t.ticket_type,
            "status": t.status,
            "prioritaet": t.priority_rank,
            "betreff": t.subject,
            "melder": t.requester_name,
            "created_at": _iso(t.created_at),
            "updated_at": _iso(t.updated_at),
        }
        for t, slug in rows
    ]


def _schichtbuch_rows(
    db: Session,
    *,
    plant_slug: str,
    from_dt: datetime | None,
    to_dt: datetime | None,
    limit: int,
) -> list[dict]:
    stmt = select(ShiftEntry, Plant.slug).join(Plant, Plant.id == ShiftEntry.plant_id).where(Plant.slug == plant_slug)
    if from_dt is not None:
        stmt = stmt.where(ShiftEntry.created_at >= from_dt)
    if to_dt is not None:
        stmt = stmt.where(ShiftEntry.created_at <= to_dt)

    rows = db.execute(stmt.order_by(ShiftEntry.created_at.desc()).limit(max(1, min(limit, 2000)))).all()
    return [
        {
            "row_type": "schichtbuch",
            "entry_id": e.id,
            "plant_slug": slug,
            "status": e.status,
            "betreff": e.subject,
            "autor": e.author_name,
            "created_at": _iso(e.created_at),
            "updated_at": _iso(e.updated_at),
        }
        for e, slug in rows
    ]


def build_report_dataset(
    db: Session,
    *,
    plant_slug: str,
    from_dt: datetime | None = None,
    to_dt: datetime | None = None,
    department: str | None = None,
    ticket_id: int | None = None,
    report_kind: str | None = "tickets",
    limit: int = 300,
) -> dict:
    kind = normalize_report_kind(report_kind)
    max_rows = max(1, min(limit, 2000))
    normalized_department = (department or "").strip() or None

    ticket_rows: list[dict] = []
    entry_rows: list[dict] = []
    if kind in {"tickets", "kombiniert"}:
        ticket_rows = _ticket_rows(
            db,
            plant_slug=plant_slug,
            from_dt=from_dt,
            to_dt=to_dt,
            department=normalized_department,
            ticket_id=ticket_id,
            limit=max_rows,
        )
    if kind in {"schichtbuch", "kombiniert"}:
        entry_rows = _schichtbuch_rows(
            db,
            plant_slug=plant_slug,
            from_dt=from_dt,
            to_dt=to_dt,
            limit=max_rows,
        )

    if kind == "tickets":
        rows = ticket_rows
    elif kind == "schichtbuch":
        rows = entry_rows
    else:
        rows = sorted(ticket_rows + entry_rows, key=lambda r: str(r.get("created_at") or ""), reverse=True)

    rows = rows[:max_rows]
    columns: list[str] = []
    for row in rows:
        for key in row.keys():
            if key not in columns:
                columns.append(key)

    ticket_status_counts: dict[str, int] = {}
    for row in ticket_rows:
        status = str(row.get("status") or "UNBEKANNT")
        ticket_status_counts[status] = ticket_status_counts.get(status, 0) + 1

    summary = {
        "report_kind": kind,
        "plant_slug": plant_slug,
        "from": _iso(from_dt),
        "to": _iso(to_dt),
        "bereich": normalized_department,
        "ticket_id": ticket_id,
        "total_rows": len(rows),
        "ticket_rows": len(ticket_rows),
        "schichtbuch_rows": len(entry_rows),
        "offene_tickets": sum(1 for row in ticket_rows if str(row.get("status") or "") not in CLOSED_TICKET_STATUSES),
        "geschlossene_tickets": sum(1 for row in ticket_rows if str(row.get("status") or "") in CLOSED_TICKET_STATUSES),
        "ticket_status_counts": ticket_status_counts,
        "generated_at": _iso(datetime.now(timezone.utc)),
    }
    return {"summary": summary, "columns": columns, "rows": rows}


def _render_csv(dataset: dict) -> bytes:
    rows = dataset.get("rows") if isinstance(dataset.get("rows"), list) else []
    columns = dataset.get("columns") if isinstance(dataset.get("columns"), list) else []
    columns = [str(c) for c in columns]
    if not columns and rows:
        columns = [str(k) for k in rows[0].keys()]

    output = io.StringIO()
    if not columns:
        writer = csv.writer(output)
        writer.writerow(["info"])
        writer.writerow(["Keine Daten fuer den gewaehlten Filter"])
        return output.getvalue().encode("utf-8")

    writer = csv.DictWriter(output, fieldnames=columns)
    writer.writeheader()
    for row in rows:
        writer.writerow({k: row.get(k) for k in columns})
    return output.getvalue().encode("utf-8")


def _render_json(dataset: dict) -> bytes:
    return json.dumps(dataset, ensure_ascii=False, indent=2).encode("utf-8")


def _render_xml(dataset: dict) -> bytes:
    root = Element("report")
    summary_node = SubElement(root, "summary")
    summary = dataset.get("summary", {})
    if isinstance(summary, dict):
        for key, value in summary.items():
            item = SubElement(summary_node, str(key))
            item.text = json.dumps(value, ensure_ascii=False) if isinstance(value, (dict, list)) else str(value)

    rows_node = SubElement(root, "rows")
    rows = dataset.get("rows") if isinstance(dataset.get("rows"), list) else []
    for row in rows:
        row_node = SubElement(rows_node, "row")
        for key, value in row.items():
            item = SubElement(row_node, str(key))
            item.text = "" if value is None else str(value)
    return tostring(root)


def _as_text(value: object) -> str:
    if value is None:
        return ""
    if isinstance(value, (dict, list)):
        return json.dumps(value, ensure_ascii=False)
    return str(value)


def _dataset_rows_and_columns(dataset: dict) -> tuple[list[dict], list[str]]:
    rows = dataset.get("rows") if isinstance(dataset.get("rows"), list) else []
    columns = dataset.get("columns") if isinstance(dataset.get("columns"), list) else []
    norm_columns = [str(c) for c in columns]
    if not norm_columns and rows:
        norm_columns = [str(k) for k in rows[0].keys()]
    return rows, norm_columns


def _render_xlsx(dataset: dict) -> bytes:
    rows, columns = _dataset_rows_and_columns(dataset)
    summary = dataset.get("summary") if isinstance(dataset.get("summary"), dict) else {}

    wb = Workbook()
    ws = wb.active
    ws.title = "Report"

    ws["A1"] = "Anlagen Report"
    ws["A1"].font = Font(bold=True, size=14)

    row_idx = 3
    for key, value in summary.items():
        ws.cell(row=row_idx, column=1, value=str(key))
        ws.cell(row=row_idx, column=2, value=_as_text(value))
        row_idx += 1

    row_idx += 1
    if columns:
        for col_idx, col_name in enumerate(columns, start=1):
            cell = ws.cell(row=row_idx, column=col_idx, value=col_name)
            cell.font = Font(bold=True)
        row_idx += 1
        for row in rows:
            for col_idx, col_name in enumerate(columns, start=1):
                ws.cell(row=row_idx, column=col_idx, value=_as_text(row.get(col_name)))
            row_idx += 1

        for col_idx, col_name in enumerate(columns, start=1):
            sample_values = [_as_text(r.get(col_name)) for r in rows[:60]]
            max_len = max([len(col_name)] + [len(v) for v in sample_values] + [8])
            ws.column_dimensions[ws.cell(row=1, column=col_idx).column_letter].width = min(50, max_len + 2)

    output = io.BytesIO()
    wb.save(output)
    return output.getvalue()


def _render_docx(dataset: dict) -> bytes:
    rows, columns = _dataset_rows_and_columns(dataset)
    summary = dataset.get("summary") if isinstance(dataset.get("summary"), dict) else {}

    doc = Document()
    doc.add_heading("Anlagen Report", level=1)

    for key, value in summary.items():
        doc.add_paragraph(f"{key}: {_as_text(value)}")

    if columns:
        doc.add_heading("Daten", level=2)
        table = doc.add_table(rows=1, cols=len(columns))
        table.style = "Table Grid"
        for idx, col_name in enumerate(columns):
            table.rows[0].cells[idx].text = col_name
        for row in rows[:500]:
            cells = table.add_row().cells
            for idx, col_name in enumerate(columns):
                cells[idx].text = _as_text(row.get(col_name))
    else:
        doc.add_paragraph("Keine Daten fuer den gewaehlten Filter.")

    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()


def _render_pdf(dataset: dict) -> bytes:
    rows, columns = _dataset_rows_and_columns(dataset)
    summary = dataset.get("summary") if isinstance(dataset.get("summary"), dict) else {}

    buffer = io.BytesIO()
    page_width, page_height = landscape(A4)
    doc = canvas.Canvas(buffer, pagesize=(page_width, page_height))

    margin_x = 32
    y = page_height - 34

    def draw_line(text: str, bold: bool = False) -> None:
        nonlocal y
        if y < 28:
            doc.showPage()
            y = page_height - 34
        doc.setFont("Helvetica-Bold" if bold else "Helvetica", 10)
        doc.drawString(margin_x, y, text[:140])
        y -= 13

    draw_line("Anlagen Report", bold=True)
    draw_line("")
    for key, value in summary.items():
        draw_line(f"{key}: {_as_text(value)}")

    draw_line("")
    if not columns:
        draw_line("Keine Daten fuer den gewaehlten Filter.")
    else:
        max_cols = min(len(columns), 9)
        selected_cols = columns[:max_cols]
        col_width = (page_width - margin_x * 2) / max_cols

        def draw_row(values: list[str], bold: bool = False) -> None:
            nonlocal y
            if y < 28:
                doc.showPage()
                y = page_height - 34
            doc.setFont("Helvetica-Bold" if bold else "Helvetica", 8)
            for idx, value in enumerate(values):
                x = margin_x + idx * col_width
                doc.drawString(x, y, value[:30])
            y -= 11

        draw_row(selected_cols, bold=True)
        for row in rows[:320]:
            draw_row([_as_text(row.get(col)) for col in selected_cols])

    doc.showPage()
    doc.save()
    return buffer.getvalue()


def render_format(fmt: str, dataset: dict) -> bytes:
    fmt = fmt.lower()
    if fmt == "csv":
        return _render_csv(dataset)
    if fmt == "json":
        return _render_json(dataset)
    if fmt == "xml":
        return _render_xml(dataset)
    if fmt == "xlsx":
        return _render_xlsx(dataset)
    if fmt == "pdf":
        return _render_pdf(dataset)
    if fmt == "docx":
        return _render_docx(dataset)
    raise ValueError(f"unsupported format {fmt}")


def run_report(db: Session, run: ReportRun, formats: list[str], context: dict | None = None) -> list[ReportArtifact]:
    settings = get_settings()
    settings.reports_dir.mkdir(parents=True, exist_ok=True)

    run.status = "running"
    run.started_at = datetime.now(timezone.utc)
    db.flush()

    context = context or {}
    dataset = build_report_dataset(
        db,
        plant_slug=str(context.get("plant_slug") or run.plant_slug),
        from_dt=context.get("from_dt") or run.range_from,
        to_dt=context.get("to_dt") or run.range_to,
        department=context.get("department"),
        ticket_id=context.get("ticket_id"),
        report_kind=context.get("report_kind"),
        limit=int(context.get("limit") or 500),
    )
    artifacts: list[ReportArtifact] = []

    for fmt in formats:
        fmt_lower = fmt.lower()
        if fmt_lower not in MIME_BY_FORMAT:
            raise ValueError(f"unsupported format {fmt_lower}")
        payload = render_format(fmt_lower, dataset)
        file_name = f"report_{run.id}_{fmt_lower}.{fmt_lower}"
        path = settings.reports_dir / file_name
        path.write_bytes(payload)
        sha = hashlib.sha256(payload).hexdigest()

        artifact = ReportArtifact(
            report_run_id=run.id,
            format=fmt_lower,
            mime_type=MIME_BY_FORMAT[fmt_lower],
            path=str(path),
            size_bytes=len(payload),
            sha256=sha,
        )
        db.add(artifact)
        artifacts.append(artifact)

    run.status = "done"
    run.finished_at = datetime.now(timezone.utc)
    return artifacts
